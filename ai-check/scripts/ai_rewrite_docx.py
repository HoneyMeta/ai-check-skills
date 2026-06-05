#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Apply user-confirmed AI-check replacements to a new DOCX file.

The script does not call AI. It only validates and applies replacement JSON
created by the assistant after the user has approved the changes.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable


@dataclass
class Replacement:
    sourceText: str
    replacementText: str
    chunkId: str = ""
    reason: str = ""
    occurrence: int | None = None


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        json.dump(data, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def normalize_replacements(payload: Any) -> list[Replacement]:
    raw = payload.get("replacements", payload) if isinstance(payload, dict) else payload
    if not isinstance(raw, list):
        raise ValueError("replacements JSON must be a list or an object with a replacements list")
    replacements: list[Replacement] = []
    for index, item in enumerate(raw, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"replacement #{index} must be an object")
        source = str(item.get("sourceText") or item.get("source") or "").strip()
        replacement = str(item.get("replacementText") or item.get("replacement") or "").strip()
        if not source:
            raise ValueError(f"replacement #{index} missing sourceText")
        if not replacement:
            raise ValueError(f"replacement #{index} missing replacementText")
        occurrence = item.get("occurrence")
        replacements.append(
            Replacement(
                sourceText=source,
                replacementText=replacement,
                chunkId=str(item.get("chunkId") or "").strip(),
                reason=str(item.get("reason") or "").strip(),
                occurrence=int(occurrence) if occurrence is not None and str(occurrence).strip() else None,
            )
        )
    return replacements


def normalize_for_match(text: str) -> str:
    return re.sub(r"(?:\s|\u00a0|\u200b|\u200c|\u200d|\u2060|\u2063)+", "", text or "")


def citation_marks(text: str) -> set[str]:
    return set(re.findall(r"\[\d+(?:[-,，、]\d+)*\]", text or ""))


def validate_replacement_citations(item: Replacement) -> list[str]:
    required = citation_marks(item.sourceText)
    current = citation_marks(item.replacementText)
    missing = sorted(required - current)
    return missing


def iter_all_paragraphs(document: Any) -> Iterable[Any]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def find_matches(document: Any, item: Replacement) -> list[dict[str, Any]]:
    source = item.sourceText
    normalized_source = normalize_for_match(source)
    matches: list[dict[str, Any]] = []
    for paragraph_index, paragraph in enumerate(iter_all_paragraphs(document)):
        text = paragraph.text or ""
        if not text.strip():
            continue
        direct = source in text
        normalized = normalized_source and normalized_source in normalize_for_match(text)
        if direct or normalized:
            matches.append(
                {
                    "paragraphIndex": paragraph_index,
                    "direct": direct,
                    "paragraphText": text,
                }
            )
    return matches


def replace_normalized_once(paragraph_text: str, source_text: str, replacement_text: str) -> tuple[str, bool]:
    if source_text in paragraph_text:
        return paragraph_text.replace(source_text, replacement_text, 1), True

    normalized_source = normalize_for_match(source_text)
    if not normalized_source:
        return paragraph_text, False

    compact_to_original: list[int] = []
    compact_chars: list[str] = []
    for index, char in enumerate(paragraph_text):
        if re.match(r"(?:\s|\u00a0|\u200b|\u200c|\u200d|\u2060|\u2063)", char):
            continue
        compact_to_original.append(index)
        compact_chars.append(char)
    compact = "".join(compact_chars)
    start = compact.find(normalized_source)
    if start < 0:
        return paragraph_text, False
    end = start + len(normalized_source) - 1
    original_start = compact_to_original[start]
    original_end = compact_to_original[end] + 1
    return paragraph_text[:original_start] + replacement_text + paragraph_text[original_end:], True


def dry_run(source_docx: Path, replacements_path: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("python-docx is required: python -m pip install python-docx") from exc

    document = Document(str(source_docx))
    replacements = normalize_replacements(load_json(replacements_path))
    items = []
    for item in replacements:
        matches = find_matches(document, item)
        missing_citations = validate_replacement_citations(item)
        items.append(
            {
                "chunkId": item.chunkId,
                "matchCount": len(matches),
                "missingCitations": missing_citations,
                "willApply": len(matches) > 0 and not missing_citations,
                "matches": matches[:5],
            }
        )
    return {
        "sourcePath": str(source_docx.resolve()),
        "replacementPath": str(replacements_path.resolve()),
        "total": len(items),
        "applicable": sum(1 for item in items if item["willApply"]),
        "items": items,
    }


def apply_replacements(source_docx: Path, replacements_path: Path, output_docx: Path, dry_run_only: bool = False) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("python-docx is required: python -m pip install python-docx") from exc

    preview = dry_run(source_docx, replacements_path)
    if dry_run_only:
        return preview

    replacements = normalize_replacements(load_json(replacements_path))
    skipped = []
    applied = []
    temp_docx = output_docx
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source_docx, temp_docx)
    document = Document(str(temp_docx))

    for item in replacements:
        missing = validate_replacement_citations(item)
        if missing:
            skipped.append({"chunkId": item.chunkId, "reason": f"missing citation marks: {', '.join(missing)}"})
            continue

        matches = find_matches(document, item)
        if not matches:
            skipped.append({"chunkId": item.chunkId, "reason": "sourceText not found"})
            continue
        if item.occurrence is not None and (item.occurrence < 1 or item.occurrence > len(matches)):
            skipped.append({"chunkId": item.chunkId, "reason": "occurrence out of range"})
            continue
        if item.occurrence is None and len(matches) > 1:
            skipped.append({"chunkId": item.chunkId, "reason": "sourceText matched multiple paragraphs; set occurrence to apply deterministically"})
            continue

        target_match = matches[(item.occurrence or 1) - 1]
        target_index = target_match["paragraphIndex"]
        for paragraph_index, paragraph in enumerate(iter_all_paragraphs(document)):
            if paragraph_index != target_index:
                continue
            updated, ok = replace_normalized_once(paragraph.text or "", item.sourceText, item.replacementText)
            if not ok:
                skipped.append({"chunkId": item.chunkId, "reason": "normalized replacement failed"})
                break
            paragraph.text = updated
            applied.append({"chunkId": item.chunkId, "paragraphIndex": paragraph_index})
            break

    document.save(str(output_docx))
    return {
        "sourcePath": str(source_docx.resolve()),
        "outputPath": str(output_docx.resolve()),
        "replacementPath": str(replacements_path.resolve()),
        "appliedCount": len(applied),
        "skippedCount": len(skipped),
        "applied": applied,
        "skipped": skipped,
        "note": "Paragraph formatting may be simplified for paragraphs whose text was replaced. Review the output DOCX before final use.",
    }


def command_dry_run(args: argparse.Namespace) -> int:
    output = dry_run(Path(args.docx), Path(args.replacements_json))
    if args.out_json:
        write_json(Path(args.out_json), output)
    print(json.dumps(output, ensure_ascii=False, indent=2))
    return 0


def command_apply(args: argparse.Namespace) -> int:
    output = apply_replacements(Path(args.docx), Path(args.replacements_json), Path(args.out), dry_run_only=False)
    if args.out_json:
        write_json(Path(args.out_json), output)
    print(json.dumps(output, ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Apply user-confirmed AI-check replacements to a new DOCX")
    sub = parser.add_subparsers(dest="command", required=True)

    dry = sub.add_parser("dry-run", help="validate replacement matching without writing DOCX")
    dry.add_argument("docx")
    dry.add_argument("replacements_json")
    dry.add_argument("--out-json")
    dry.set_defaults(func=command_dry_run)

    apply = sub.add_parser("apply", help="apply replacements to a new DOCX")
    apply.add_argument("docx")
    apply.add_argument("replacements_json")
    apply.add_argument("--out", required=True, help="new DOCX path")
    apply.add_argument("--out-json", help="write apply summary JSON")
    apply.set_defaults(func=command_apply)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return int(args.func(args))
    except Exception as exc:
        print(f"AI-check rewrite error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())

