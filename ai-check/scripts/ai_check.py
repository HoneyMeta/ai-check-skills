#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""AI-check deterministic DOCX scanner and report generator.

This script deliberately does not call any AI model and does not calculate final
AI-writing probabilities. It extracts DOCX text, builds chunks, scans
explainable trace features, and renders reports from assistant-provided results.
"""

from __future__ import annotations

import argparse
import html
import json
import math
import os
import re
import sys
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Literal


SectionType = Literal[
    "abstract",
    "keywords",
    "chapter_title",
    "body",
    "conclusion",
    "references",
    "other",
]
Importance = Literal["high", "medium", "low"]
Sensitivity = Literal["high", "medium", "low"]
Verdict = Literal["ai", "human", "suspicious"]


MAX_CHARS_PER_CHUNK = 900
REPORT_MIN_VALUE = 0.30
REPORT_FRAGMENT_MAX_CHARS = 360
REPORT_FRAGMENT_MIN_CHARS = 24
OFFICIAL_SITE = "https://honeymeta.com/easyidea/"


TEMPLATE_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("首先其次最后模板", re.compile(r"首先[\s\S]{0,120}其次[\s\S]{0,120}(?:最后|最终)")),
    ("此外同时另外堆叠", re.compile(r"此外[\s\S]{0,120}同时[\s\S]{0,120}另外")),
    ("一方面另一方面对称句", re.compile(r"一方面[\s\S]{0,120}另一方面")),
    ("不仅而且对称句", re.compile(r"不仅[\s\S]{0,120}而且")),
    ("综上所述", re.compile(r"综上所述")),
    ("由此可见", re.compile(r"由此可见")),
    ("值得注意的是", re.compile(r"值得注意的是")),
    ("不可否认", re.compile(r"不可否认")),
    ("毋庸置疑", re.compile(r"毋庸置疑")),
    ("显而易见", re.compile(r"显而易见")),
    ("众所周知", re.compile(r"众所周知")),
    ("不言而喻", re.compile(r"不言而喻")),
]

CONNECTORS = (
    "因此",
    "然而",
    "同时",
    "此外",
    "另外",
    "由此可见",
    "综上",
    "进一步",
    "相较之下",
    "具体而言",
    "首先",
    "其次",
    "最后",
    "一方面",
    "另一方面",
)

GENERIC_CLAIM_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("重要意义套话", re.compile(r"(?:具有|具备).{0,12}(?:重要|重大).{0,12}(?:意义|价值)")),
    ("广阔前景套话", re.compile(r"(?:广阔|良好).{0,8}(?:应用)?前景")),
    ("理论实践价值套话", re.compile(r"理论价值.{0,12}现实意义|现实意义.{0,12}理论价值")),
    ("高效可靠可信泛化目标", re.compile(r"(?:高效|可靠|可信|智能化|实用化).{0,18}(?:系统|框架|方案)")),
    ("深入全面研究", re.compile(r"(?:深入|全面).{0,8}研究")),
    ("一系列问题挑战", re.compile(r"一系列.{0,8}(?:问题|挑战)")),
]

EVIDENCE_MARKER_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("引用编号", re.compile(r"\[\d+(?:[-,，、]\d+)*\]")),
    (
        "实验指标",
        re.compile(r"\d+(?:\.\d+)?\s*(?:%|G\s*FLOPs?|FLOPs?|MB|GB|ms|s|Hz|kHz|MHz|dB|mAP|AUC|F1)", re.I),
    ),
    (
        "模型或方法名",
        re.compile(r"\b(?:CNN|RNN|LSTM|GRU|Transformer|BSCQT|DCAM|DCAMNet|AB-CQT|DFAM|DFAMNet|Grad-CAM|XGBoost|CQT)\b"),
    ),
    ("实验结果表述", re.compile(r"实验结果(?:表明|显示|验证)|准确率|计算复杂度|消融|对比实验|数据集")),
    (
        "限定或取舍",
        re.compile(r"在.{0,18}(?:条件|场景|约束|限制)下|针对.{0,24}(?:问题|场景|任务)|本文(?:聚焦|采用|选取|限定)"),
    ),
]


@dataclass
class TraceFeatures:
    riskSignals: list[str]
    mitigatingSignals: list[str]
    templatePhrases: list[str]
    connectorCount: int
    connectorDensity: float
    sentenceCount: int
    sentenceLengthVariance: float | None
    uniformSentenceRhythm: bool
    genericClaimCount: int
    evidenceMarkers: list[str]


@dataclass
class Chunk:
    chunkId: str
    sectionType: SectionType
    sectionTitle: str | None
    text: str
    charCount: int
    importance: Importance
    traceFeatures: TraceFeatures


@dataclass
class Result:
    chunkId: str
    sensitivity: Sensitivity
    aigcValue: float
    verdict: Verdict
    reason: str
    signals: list[str]


def unique_limited(values: Iterable[str], limit: int = 8) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        item = str(value).strip()
        if not item or item in seen:
            continue
        seen.add(item)
        out.append(item)
        if len(out) >= limit:
            break
    return out


def count_chars(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def normalize_paragraphs(text: str) -> list[str]:
    return [
        re.sub(r"\s+", " ", line).strip()
        for line in re.split(r"\n+", (text or "").replace("\r\n", "\n").replace("\r", "\n"))
        if re.sub(r"\s+", " ", line).strip()
    ]


def split_sentences(text: str) -> list[str]:
    return [
        re.sub(r"\s+", " ", item).strip()
        for item in re.split(r"(?<=[。！？!?；;])\s*", text or "")
        if count_chars(item) > 5
    ]


def pattern_count(text: str, pattern: re.Pattern[str]) -> int:
    return len(pattern.findall(text or ""))


def scan_trace_features(text: str) -> TraceFeatures:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    sentences = split_sentences(normalized)
    sentence_lengths = [count_chars(sentence) for sentence in sentences]
    variance: float | None = None
    if sentence_lengths:
        avg = sum(sentence_lengths) / len(sentence_lengths)
        variance = sum((length - avg) ** 2 for length in sentence_lengths) / len(sentence_lengths)
    uniform = len(sentences) >= 4 and variance is not None and variance < 120

    template_phrases = unique_limited(label for label, pattern in TEMPLATE_PATTERNS if pattern.search(normalized))
    connector_count = sum(normalized.count(connector) for connector in CONNECTORS)
    connector_density = round(connector_count / len(sentences), 3) if sentences else 0.0

    generic_labels: list[str] = []
    generic_count = 0
    for label, pattern in GENERIC_CLAIM_PATTERNS:
        count = pattern_count(normalized, pattern)
        if count:
            generic_count += count
            generic_labels.append(label)

    evidence_markers = unique_limited(
        label for label, pattern in EVIDENCE_MARKER_PATTERNS if pattern.search(normalized)
    )

    risk_signals = [
        *(f"模板表达:{label}" for label in template_phrases),
        "句式节奏过于均匀" if uniform else "",
        "连接词密集" if connector_density >= 0.45 and connector_count >= 3 else "",
        f"泛化套话:{'/'.join(unique_limited(generic_labels, 3))}" if generic_count else "",
        "贡献罗列结构"
        if re.search(r"主要研究工作如下|具体研究内容如下|本文(?:提出|构建|设计).{0,80}(?:方法|框架).{0,80}(?:实验结果|结果表明)", normalized)
        else "",
        "目标表述泛化" if re.search(r"(?:旨在|力求|致力于).{0,40}(?:高效|可靠|可信|智能化|实用化)", normalized) else "",
    ]

    mitigating = [
        *(f"具体证据:{label}" for label in evidence_markers),
        "领域方法细节" if re.search(r"\b(?:BSCQT|DCAM|DCAMNet|AB-CQT|DFAM|DFAMNet|Grad-CAM|XGBoost)\b", normalized) else "",
        "定量结果支撑" if re.search(r"\d+(?:\.\d+)?\s*%", normalized) and re.search(r"准确率|提升|降低|复杂度|FLOPs?", normalized, re.I) else "",
        "限制条件说明" if re.search(r"(?:但|然而|不过).{0,40}(?:限制|不足|难以|约束|挑战)", normalized) else "",
    ]

    return TraceFeatures(
        riskSignals=unique_limited(risk_signals),
        mitigatingSignals=unique_limited(mitigating),
        templatePhrases=template_phrases,
        connectorCount=connector_count,
        connectorDensity=connector_density,
        sentenceCount=len(sentences),
        sentenceLengthVariance=round(variance, 3) if variance is not None else None,
        uniformSentenceRhythm=uniform,
        genericClaimCount=generic_count,
        evidenceMarkers=evidence_markers,
    )


def infer_section_type(text: str) -> SectionType:
    compact = re.sub(r"\s+", "", text or "")
    if re.fullmatch(r"目录|目次|TableofContents", compact, re.I):
        return "other"
    if re.match(r"^(摘要|中文摘要|Abstract)([:：].*)?$", compact, re.I):
        return "abstract"
    if re.match(r"^(关键词|关键字|Keywords?)[:：]?", compact, re.I):
        return "keywords"
    if re.match(r"^(结论|总结|结束语|展望|Conclusion)([:：].*)?$", compact, re.I):
        return "conclusion"
    if re.match(r"^(参考文献|References)$", compact, re.I):
        return "references"
    if re.match(r"^(第[一二三四五六七八九十百千万\d]+[章节部分篇]|[0-9]+(?:\.[0-9]+){0,4}\s*[\u4e00-\u9fffA-Za-z])", text or ""):
        return "chapter_title"
    return "body"


def is_standalone_heading(text: str, section_type: SectionType) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if section_type == "abstract":
        return bool(re.fullmatch(r"摘要|中文摘要|Abstract", compact, re.I))
    if section_type == "keywords":
        return bool(re.fullmatch(r"关键词|关键字|Keywords?[:：]?", compact, re.I))
    if section_type == "conclusion":
        return bool(re.fullmatch(r"结论|总结|结束语|展望|Conclusion", compact, re.I))
    if section_type == "references":
        return bool(re.fullmatch(r"参考文献|References", compact, re.I))
    return False


def importance_for(section_type: SectionType) -> Importance:
    if section_type in ("abstract", "conclusion"):
        return "high"
    if section_type in ("references", "keywords", "chapter_title"):
        return "low"
    return "medium"


def is_detectable(section_type: SectionType) -> bool:
    return section_type in ("abstract", "body", "conclusion")


def is_toc_entry(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return False
    if re.fullmatch(r"目录|目次|TableofContents", compact, re.I):
        return True
    if re.match(r"^(目录|目次)", compact) and re.search(r"(?:摘要|Abstract|第[一二三四五六七八九十百千万\d]+章|[0-9]+(?:\.[0-9]+){1,4})", compact, re.I):
        return True
    return bool(
        re.match(
            r"^(摘要|Abstract|ABSTRACT|第[一二三四五六七八九十百千万\d]+章|[0-9]+(?:\.[0-9]+){1,4}).{0,80}(?:[ivxlcdm]+|\d+)$",
            compact,
            re.I,
        )
    )


def is_cover_or_metadata(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact or len(compact) <= 2:
        return True
    if re.search(r"原创性声明|授权声明|使用授权|版权|学位论文|毕业论文|学校代码|分类号|密级|学号|作者签名|导师签名|指导教师|学院|专业|研究方向|答辩委员会|论文题目|英文题目", compact):
        return True
    if len(compact) <= 260 and not re.search(r"[。！？!?；;]", text or ""):
        return True
    return bool(re.fullmatch(r"摘要|中文摘要|Abstract|关键词|关键字|Keywords?|目录|目次|参考文献|References", compact, re.I))


def strip_section_label(text: str, section_type: SectionType) -> str:
    normalized = (text or "").strip()
    if section_type == "abstract":
        return re.sub(r"^(?:摘\s*要|中文摘要|Abstract)\s*[:：]?\s*", "", normalized, flags=re.I).strip() or normalized
    if section_type == "conclusion":
        return re.sub(r"^(?:结论|总结|结束语|展望|Conclusion)\s*[:：]?\s*", "", normalized, flags=re.I).strip() or normalized
    return normalized


def is_non_content_text(text: str, section_type: SectionType | None = None) -> bool:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    compact = re.sub(r"\s+", "", normalized)
    if not compact:
        return True
    if section_type and not is_detectable(section_type):
        return True
    if is_toc_entry(normalized):
        return True
    if re.match(r"^(?:关键词|关键字|Keywords?)[:：]?", compact, re.I):
        return True
    if re.fullmatch(r"参考文献|References", compact, re.I):
        return True
    if re.match(r"^(?:第[一二三四五六七八九十百千万\d]+[章节部分篇]|[0-9]+(?:\.[0-9]+){0,4})", compact):
        return True
    if count_chars(normalized) <= 60 and re.match(r"^(?:\([一二三四五六七八九十\d]+\)|（[一二三四五六七八九十\d]+）|[一二三四五六七八九十\d]+[、.．])", normalized) and not re.search(r"[。！？!?；;，,：:]", normalized):
        return True
    if re.match(r"^(?:图|表)\s*\d+(?:[.-]\d+)*\s*", normalized):
        return True
    return False


def split_paragraph(paragraph: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> list[str]:
    normalized = paragraph.strip()
    if count_chars(normalized) <= max_chars:
        return [normalized]
    pieces: list[str] = []
    current = ""
    for sentence in re.split(r"(?<=[。！？!?；;])\s*", normalized):
        sentence = sentence.strip()
        if not sentence:
            continue
        if current and count_chars(current + sentence) > max_chars:
            pieces.append(current.strip())
            current = sentence
        else:
            current += sentence
    if current.strip():
        pieces.append(current.strip())
    return pieces or [normalized[:max_chars]]


def compact_chunks(chunks: list[Chunk], max_chars: int = MAX_CHARS_PER_CHUNK) -> list[Chunk]:
    compacted: list[Chunk] = []
    pending: dict[str, Any] | None = None

    def flush() -> None:
        nonlocal pending
        if not pending:
            return
        text = str(pending["text"])
        compacted.append(
            Chunk(
                chunkId=f"chunk-{len(compacted) + 1:04d}",
                sectionType=pending["sectionType"],
                sectionTitle=pending.get("sectionTitle"),
                text=text,
                charCount=count_chars(text),
                importance=pending["importance"],
                traceFeatures=scan_trace_features(text),
            )
        )
        pending = None

    for chunk in chunks:
        if not is_detectable(chunk.sectionType):
            flush()
            continue
        if (
            pending
            and pending["sectionType"] == chunk.sectionType
            and pending["importance"] == chunk.importance
            and (pending.get("sectionTitle") or "") == (chunk.sectionTitle or "")
            and count_chars(str(pending["text"])) + chunk.charCount <= max_chars
        ):
            pending["text"] = f"{pending['text']}\n{chunk.text}"
            continue
        flush()
        pending = {
            "sectionType": chunk.sectionType,
            "sectionTitle": chunk.sectionTitle,
            "text": chunk.text,
            "importance": chunk.importance,
        }
    flush()
    return compacted


def build_chunks(text: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> list[Chunk]:
    paragraphs = normalize_paragraphs(text)
    chunks: list[Chunk] = []
    current_title = ""
    current_type: SectionType = "body"
    expect_keyword_payload = False
    before_main = True
    in_toc = False

    for paragraph in paragraphs:
        inferred = infer_section_type(paragraph)
        compact = re.sub(r"\s+", "", paragraph)

        if re.fullmatch(r"目录|目次|TableofContents", compact, re.I):
            in_toc = True
            expect_keyword_payload = False
            current_type = "other"
            continue

        if inferred == "chapter_title":
            if in_toc and is_toc_entry(paragraph):
                continue
            in_toc = False
            before_main = False
            current_title = paragraph[:120]
            current_type = "conclusion" if re.search(r"结论|总结|展望|Conclusion", paragraph, re.I) else "body"
            expect_keyword_payload = False
            continue

        section_type = "keywords" if inferred == "body" and expect_keyword_payload else inferred if inferred != "body" else current_type
        standalone = is_standalone_heading(paragraph, inferred)
        if standalone and inferred in ("abstract", "keywords", "conclusion", "references"):
            current_title = paragraph[:120]
            expect_keyword_payload = inferred == "keywords"
            current_type = "body" if inferred == "keywords" else inferred
            if inferred in ("abstract", "conclusion", "references"):
                in_toc = False
            if inferred in ("abstract", "conclusion"):
                before_main = False
            continue
        if inferred != "keywords" or not standalone:
            expect_keyword_payload = False

        if in_toc and (is_toc_entry(paragraph) or count_chars(paragraph) < 80):
            continue
        if not is_detectable(section_type):
            continue
        if before_main and is_cover_or_metadata(paragraph):
            continue

        stripped = strip_section_label(paragraph, section_type)
        if is_non_content_text(stripped, section_type):
            continue
        before_main = False

        for piece in split_paragraph(stripped, max_chars=max_chars):
            char_count = count_chars(piece)
            if char_count < REPORT_FRAGMENT_MIN_CHARS:
                continue
            chunks.append(
                Chunk(
                    chunkId=f"chunk-{len(chunks) + 1:04d}",
                    sectionType=section_type,
                    sectionTitle=current_title or None,
                    text=piece,
                    charCount=char_count,
                    importance=importance_for(section_type),
                    traceFeatures=scan_trace_features(piece),
                )
            )

    return compact_chunks(chunks, max_chars=max_chars)


def iter_table_text(table: Any) -> Iterable[str]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                if text:
                    yield text
            for nested in cell.tables:
                yield from iter_table_text(nested)


def read_docx_text(path: Path) -> tuple[str, dict[str, str]]:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("python-docx is required for reading DOCX files: python -m pip install python-docx") from exc

    document = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        lines.extend(iter_table_text(table))
    props = document.core_properties
    metadata = {
        "title": (props.title or "").strip(),
        "author": (props.author or "").strip(),
        "subject": (props.subject or "").strip(),
    }
    return "\n".join(lines), metadata


def is_probable_document_title(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    compact = re.sub(r"\s+", "", normalized)
    if not (4 <= count_chars(normalized) <= 80):
        return False
    if infer_section_type(normalized) != "body":
        return False
    if is_toc_entry(normalized) or is_non_content_text(normalized):
        return False
    if re.search(r"[。！？!?；;]", normalized):
        return False
    if re.search(
        r"原创性声明|授权声明|使用授权|版权|学位论文|毕业论文|学校代码|分类号|密级|学号|"
        r"作者签名|导师签名|指导教师|学院|专业|研究方向|答辩委员会|论文题目|英文题目",
        compact,
    ):
        return False
    return True


def summarize_title(path: Path, text: str, explicit_title: str | None = None, metadata_title: str | None = None) -> str:
    if explicit_title and explicit_title.strip():
        return explicit_title.strip()[:120]
    if metadata_title and metadata_title.strip():
        return metadata_title.strip()[:120]
    for line in normalize_paragraphs(text):
        inferred = infer_section_type(line)
        if inferred in ("abstract", "chapter_title", "conclusion", "references"):
            break
        if is_probable_document_title(line):
            return line[:120]
    return path.stem or "untitled"


def scan_docx(path: Path, title: str | None = None, author: str | None = None, max_chars: int = MAX_CHARS_PER_CHUNK) -> dict[str, Any]:
    text, metadata = read_docx_text(path)
    chunks = build_chunks(text, max_chars=max_chars)
    return {
        "schemaVersion": 1,
        "tool": "AI-check",
        "sourcePath": str(path.resolve()),
        "documentTitle": summarize_title(path, text, explicit_title=title, metadata_title=metadata.get("title")),
        "author": author or metadata.get("author") or "unknown",
        "createdAt": datetime.now().astimezone().isoformat(timespec="seconds"),
        "totalChars": sum(chunk.charCount for chunk in chunks),
        "chunking": {"maxCharsPerChunk": max_chars, "minReportChars": REPORT_FRAGMENT_MIN_CHARS},
        "chunks": [chunk_to_dict(chunk) for chunk in chunks],
    }


def chunk_to_dict(chunk: Chunk) -> dict[str, Any]:
    data = asdict(chunk)
    data["traceFeatures"] = asdict(chunk.traceFeatures)
    return data


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        json.dump(data, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def normalize_results(payload: Any) -> list[Result]:
    raw_results = payload.get("results", payload) if isinstance(payload, dict) else payload
    if not isinstance(raw_results, list):
        raise ValueError("results JSON must be a list or an object with a results list")

    results: list[Result] = []
    for item in raw_results:
        if not isinstance(item, dict):
            raise ValueError("each result must be an object")
        value = float(item.get("aigcValue"))
        if not math.isfinite(value) or value < 0 or value > 1:
            raise ValueError(f"aigcValue must be 0..1 for {item.get('chunkId')}")
        sensitivity = str(item.get("sensitivity") or "medium")
        verdict = str(item.get("verdict") or "suspicious")
        if sensitivity not in ("low", "medium", "high"):
            raise ValueError(f"invalid sensitivity for {item.get('chunkId')}")
        if verdict not in ("human", "suspicious", "ai"):
            raise ValueError(f"invalid verdict for {item.get('chunkId')}")
        signals = item.get("signals") or []
        if isinstance(signals, str):
            signals = [part.strip() for part in re.split(r"[,，、\n]+", signals) if part.strip()]
        results.append(
            Result(
                chunkId=str(item.get("chunkId") or "").strip(),
                sensitivity=sensitivity,  # type: ignore[arg-type]
                aigcValue=round(value, 3),
                verdict=verdict,  # type: ignore[arg-type]
                reason=str(item.get("reason") or "").strip(),
                signals=[str(signal).strip() for signal in signals if str(signal).strip()],
            )
        )
    return results


def comparable_text(value: str) -> str:
    return re.sub(r"(?:\s|\u00a0|\u200b|\u200c|\u200d|\u2060|\u2063)+", "", value or "").strip()


def build_report_fragments(text: str) -> list[str]:
    paragraphs = normalize_paragraphs(text)
    fragments: list[str] = []
    for paragraph in paragraphs:
        if is_non_content_text(paragraph):
            continue
        current = ""
        for sentence in re.split(r"(?<=[。！？!?；;])\s*", paragraph):
            normalized = sentence.strip()
            if not normalized:
                continue
            next_text = f"{current}{normalized}" if current else normalized
            if current and count_chars(next_text) > REPORT_FRAGMENT_MAX_CHARS:
                fragments.append(current)
                current = normalized
            else:
                current = next_text
        if current:
            fragments.append(current)

    fallback = [paragraph for paragraph in paragraphs if not is_non_content_text(paragraph)]
    candidates = fragments or fallback
    out: list[str] = []
    for fragment in candidates:
        out.extend(split_long_fragment(fragment))
    return [fragment for fragment in out if count_chars(fragment) >= REPORT_FRAGMENT_MIN_CHARS]


def split_long_fragment(fragment: str) -> list[str]:
    normalized = re.sub(r"\s+", " ", fragment or "").strip()
    if not normalized:
        return []
    if count_chars(normalized) <= REPORT_FRAGMENT_MAX_CHARS:
        return [normalized]

    pieces: list[str] = []
    rest = normalized
    while count_chars(rest) > REPORT_FRAGMENT_MAX_CHARS:
        hard_limit = min(len(rest), REPORT_FRAGMENT_MAX_CHARS)
        window = rest[:hard_limit]
        min_boundary = max(48, hard_limit // 2)
        split_at = -1
        for token in ("。", "！", "？", "；", ";", "，", ",", "、", " "):
            index = window.rfind(token)
            if index >= min_boundary:
                split_at = max(split_at, index + len(token))
        if split_at <= 0:
            split_at = hard_limit
        piece = rest[:split_at].strip()
        if piece:
            pieces.append(piece)
        rest = rest[split_at:].strip()
        if not rest:
            break
    if rest:
        pieces.append(rest)
    return pieces


def report_reason(result: Result) -> str:
    reason = result.reason.strip()
    signals = [signal.strip() for signal in result.signals if signal.strip()]
    if not signals:
        return reason
    return f"{reason} 信号：{'、'.join(signals)}" if reason else f"信号：{'、'.join(signals)}"


def percent(value: float) -> str:
    return f"{round(max(0.0, min(1.0, value)) * 1000) / 10:.1f}%"


def get_report_rows(scan: dict[str, Any], results: list[Result]) -> list[dict[str, Any]]:
    result_by_chunk = {result.chunkId: result for result in results}
    rows: list[dict[str, Any]] = []
    seen: set[str] = set()
    for chunk in scan.get("chunks", []):
        result = result_by_chunk.get(str(chunk.get("chunkId") or ""))
        if not result or result.aigcValue < REPORT_MIN_VALUE:
            continue
        if chunk.get("sectionType") not in ("abstract", "body", "conclusion"):
            continue
        for text in build_report_fragments(str(chunk.get("text") or "")):
            normalized = comparable_text(text)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            rows.append({"chunk": chunk, "result": result, "text": text})
    return rows


def calculate_summary(scan: dict[str, Any], rows: list[dict[str, Any]]) -> dict[str, Any]:
    total_chars = int(scan.get("totalChars") or 0)
    if total_chars <= 0:
        total_chars = sum(int(chunk.get("charCount") or count_chars(str(chunk.get("text") or ""))) for chunk in scan.get("chunks", []))
    total_chars = max(1, total_chars)
    suspicious_chars = sum(count_chars(row["text"]) for row in rows)
    weighted_score = sum(count_chars(row["text"]) * row["result"].aigcValue for row in rows)
    ai_rate = weighted_score / total_chars
    suspicious_ai_rate = suspicious_chars / total_chars
    return {
        "aiRate": round(max(0.0, min(1.0, ai_rate)), 4),
        "suspiciousAiRate": round(max(0.0, min(1.0, suspicious_ai_rate)), 4),
        "humanRate": round(max(0.0, min(1.0, 1 - ai_rate)), 4),
        "suspiciousChars": suspicious_chars,
        "totalChars": total_chars,
        "fragmentCount": len(rows),
    }


def render_html(scan: dict[str, Any], rows: list[dict[str, Any]], summary: dict[str, Any]) -> str:
    detection_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    title = str(scan.get("documentTitle") or "untitled")
    author = str(scan.get("author") or "unknown")
    distribution = "\n".join(
        f'<span class="dist {"high" if row["result"].aigcValue >= 0.75 else "medium" if row["result"].aigcValue >= 0.45 else "low"}" style="height:{max(8, round(12 + row["result"].aigcValue * 42))}px"></span>'
        for row in rows
    )
    table_rows = "\n".join(
        f"""<tr>
  <td class="original">
    <div class="original-text">{html.escape(row["text"])}</div>
    <div class="reason">检测原因：{html.escape(report_reason(row["result"]))}</div>
  </td>
  <td class="score">AIGC值：{row["result"].aigcValue:.3f}</td>
</tr>"""
        for row in rows
    )
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>AI-check / EasyIdea 论文AI率检测报告</title>
  <style>
    @page {{ size: A4; margin: 18mm; }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; color: #202124; font-family: "Microsoft YaHei", "SimSun", "Noto Sans CJK SC", Arial, sans-serif; font-size: 12px; line-height: 1.65; }}
    .page-title {{ text-align: center; font-size: 22px; font-weight: 700; margin: 2mm 0 8mm; }}
    .meta {{ display: grid; grid-template-columns: 1fr 1fr; gap: 2mm 8mm; border: 1px solid #d9d9d9; padding: 4mm 5mm; margin-bottom: 6mm; }}
    .section-title {{ background: #f2f4f7; border-left: 4px solid #4b89dc; font-size: 15px; font-weight: 700; padding: 2mm 3mm; margin: 5mm 0 3mm; }}
    .metrics {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 3mm; margin-bottom: 5mm; }}
    .metric {{ border: 1px solid #dadce0; min-height: 20mm; padding: 3mm; text-align: center; }}
    .metric strong {{ display: block; color: #d93025; font-size: 20px; line-height: 1.2; }}
    .metric span {{ color: #5f6368; }}
    .distribution {{ height: 22mm; border: 1px solid #dadce0; display: flex; align-items: flex-end; gap: 1px; padding: 2mm; margin-bottom: 5mm; }}
    .dist {{ display: inline-block; width: 5px; min-width: 3px; background: #34a853; }}
    .dist.medium {{ background: #fbbc04; }}
    .dist.high {{ background: #ea4335; }}
    table {{ width: 100%; border-collapse: collapse; table-layout: fixed; }}
    tr {{ break-inside: avoid; page-break-inside: avoid; }}
    th {{ background: #f7f7f7; font-weight: 700; }}
    th, td {{ border: 1px solid #d9d9d9; vertical-align: top; padding: 2.2mm 2.4mm; }}
    td.original {{ width: 76%; word-break: break-all; }}
    .reason {{ margin-top: 1.6mm; color: #6b7280; font-size: 10.5px; line-height: 1.45; }}
    td.score {{ width: 24%; color: #d93025; white-space: nowrap; font-weight: 700; }}
    .notes {{ margin-top: 7mm; color: #4b5563; }}
    .site-link {{ color: #1a73e8; text-decoration: underline; }}
  </style>
</head>
<body>
  <h1 class="page-title">AI-check / EasyIdea 论文AI率检测报告</h1>
  <section class="meta">
    <div>检测时间：{html.escape(detection_time)}</div>
    <div>检测文献：{html.escape(title)}</div>
    <div>作者：{html.escape(author)}</div>
    <div>检测类型：AIGC 写作检测</div>
    <div>官网：<a class="site-link" href="{OFFICIAL_SITE}">{OFFICIAL_SITE}</a></div>
    <div style="grid-column:1 / -1">检测范围：ChatGPT 讯飞星火 Gemini Kimi Claude 文心一言 通义千问 智谱AI 百川智能 360智脑 豆包 DeepSeek (包括但不限于)</div>
  </section>
  <div class="section-title">检测结果</div>
  <section class="metrics">
    <div class="metric"><strong>{percent(summary["aiRate"])}</strong><span>AI率</span></div>
    <div class="metric"><strong>{percent(summary["suspiciousAiRate"])}</strong><span>疑似AI率</span></div>
    <div class="metric"><strong>{percent(summary["humanRate"])}</strong><span>人写作率</span></div>
    <div class="metric"><strong>[{summary["suspiciousChars"]}]</strong><span>疑似AI写作字数</span></div>
    <div class="metric"><strong>[{summary["totalChars"]}]</strong><span>总字数</span></div>
  </section>
  <div class="section-title">疑似片段分布图</div>
  <section class="distribution">{distribution}</section>
  <table>
    <thead><tr><th style="width:76%">原文内容</th><th style="width:24%">疑似AI写作率</th></tr></thead>
    <tbody>{table_rows}</tbody>
  </table>
  <section class="notes">
    <p>说明：</p>
    <p>本报告由 AI-check 根据当前助手逐段判断结果生成，AIGC值与文章质量无关，仅供参考。</p>
    <p>本报告不代表知网、维普、Turnitin 或其他闭源平台检测结果。</p>
    <p>疑似AI生成段落中的“片段”为检测自动划分，与原文自然段可能不同。</p>
    <p>EasyIdea官网：<a class="site-link" href="{OFFICIAL_SITE}">{OFFICIAL_SITE}</a></p>
  </section>
</body>
</html>
"""


def find_cjk_font() -> str | None:
    candidates = [
        os.environ.get("AI_CHECK_PDF_FONT"),
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def wrap_text_by_width(text: str, max_units: int) -> list[str]:
    lines: list[str] = []
    current = ""
    units = 0
    for char in text:
        char_units = 2 if "\u4e00" <= char <= "\u9fff" else 1
        if current and units + char_units > max_units:
            lines.append(current)
            current = char
            units = char_units
        else:
            current += char
            units += char_units
    if current:
        lines.append(current)
    return lines or [""]


def render_pdf(path: Path, scan: dict[str, Any], rows: list[dict[str, Any]], summary: dict[str, Any]) -> None:
    try:
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import ParagraphStyle  # type: ignore
        from reportlab.lib.units import mm  # type: ignore
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
        from reportlab.platypus import Flowable, LongTable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("reportlab is required for PDF output: python -m pip install reportlab") from exc

    def register_pdf_font() -> str:
        cjk_font = find_cjk_font()
        if cjk_font:
            try:
                font_name = "AI_CHECK_CJK"
                pdfmetrics.registerFont(TTFont(font_name, cjk_font))
                pdfmetrics.registerFontFamily(font_name, normal=font_name, bold=font_name, italic=font_name, boldItalic=font_name)
                return font_name
            except Exception:
                pass
        fallback = "STSong-Light"
        pdfmetrics.registerFont(UnicodeCIDFont(fallback))
        return fallback

    class DistributionBars(Flowable):
        def __init__(self, values: list[float], height: float = 22 * mm) -> None:
            super().__init__()
            self.values = [max(0.0, min(1.0, float(value))) for value in values]
            self.height = height
            self.width = 0

        def wrap(self, avail_width: float, avail_height: float) -> tuple[float, float]:
            self.width = avail_width
            return avail_width, self.height

        def draw(self) -> None:
            pad = 5
            base_y = pad
            chart_h = self.height - pad * 2
            self.canv.setFillColor(colors.HexColor("#FBFCFE"))
            self.canv.setStrokeColor(colors.HexColor("#D7DEE8"))
            self.canv.roundRect(0, 0, self.width, self.height, 4, stroke=1, fill=1)
            if not self.values:
                return
            gap = 1.2
            max_bars = max(1, int((self.width - pad * 2) / 2.6))
            values = self.values[:max_bars]
            available = self.width - pad * 2 - gap * max(0, len(values) - 1)
            bar_w = max(1.4, min(5.0, available / max(1, len(values))))
            x = pad
            for value in values:
                color = "#D93025" if value >= 0.75 else "#F9AB00" if value >= 0.45 else "#34A853"
                height = max(5, value * chart_h)
                self.canv.setFillColor(colors.HexColor(color))
                self.canv.roundRect(x, base_y, bar_w, height, 1.1, stroke=0, fill=1)
                x += bar_w + gap

    font_name = register_pdf_font()
    path.parent.mkdir(parents=True, exist_ok=True)

    page_w, _page_h = A4
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=15 * mm,
        title="AI-check / EasyIdea 论文AI率检测报告",
        author="AI-check",
    )

    blue = colors.HexColor("#2563EB")
    blue_dark = colors.HexColor("#1D4ED8")
    ink = colors.HexColor("#202124")
    muted = colors.HexColor("#5F6368")
    border = colors.HexColor("#D7DEE8")
    soft = colors.HexColor("#F6F8FB")
    softer = colors.HexColor("#FBFCFE")
    red = colors.HexColor("#D93025")
    green = colors.HexColor("#188038")
    amber = colors.HexColor("#B06000")

    base = ParagraphStyle("AICheckBase", fontName=font_name, fontSize=9.2, leading=13.2, textColor=ink, alignment=TA_LEFT, wordWrap="CJK")
    title_style = ParagraphStyle("AICheckTitle", parent=base, fontSize=20, leading=26, alignment=TA_CENTER, textColor=colors.HexColor("#111827"), spaceAfter=3 * mm)
    subtitle_style = ParagraphStyle("AICheckSubtitle", parent=base, fontSize=8.8, leading=12, alignment=TA_CENTER, textColor=muted, spaceAfter=6 * mm)
    meta_style = ParagraphStyle("AICheckMeta", parent=base, fontSize=8.7, leading=12.2)
    section_style = ParagraphStyle("AICheckSection", parent=base, fontSize=11.2, leading=15, textColor=colors.HexColor("#111827"))
    metric_value_style = ParagraphStyle("AICheckMetricValue", parent=base, fontSize=15, leading=18, alignment=TA_CENTER, textColor=red)
    metric_label_style = ParagraphStyle("AICheckMetricLabel", parent=base, fontSize=7.6, leading=10, alignment=TA_CENTER, textColor=muted)
    table_header_style = ParagraphStyle("AICheckTableHeader", parent=base, fontSize=8.6, leading=11, alignment=TA_CENTER, textColor=colors.HexColor("#111827"))
    body_style = ParagraphStyle("AICheckBody", parent=base, fontSize=8.6, leading=12.8)
    reason_style = ParagraphStyle("AICheckReason", parent=base, fontSize=7.3, leading=10.2, textColor=colors.HexColor("#6B7280"))
    score_style = ParagraphStyle("AICheckScore", parent=base, fontSize=8.6, leading=12, alignment=TA_RIGHT)
    note_style = ParagraphStyle("AICheckNote", parent=base, fontSize=8.0, leading=11.5, textColor=muted)

    def safe(value: Any) -> str:
        return html.escape(str(value or "")).replace("\n", "<br/>")

    def paragraph(value: Any, style: ParagraphStyle = base) -> Paragraph:
        return Paragraph(safe(value), style)

    def color_for_value(value: float) -> Any:
        if value >= 0.75:
            return red
        if value >= 0.45:
            return amber
        return green

    def section_title(text: str) -> Table:
        table = Table([["", paragraph(text, section_style)]], colWidths=[3.5 * mm, doc.width - 3.5 * mm], hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), blue),
                    ("BACKGROUND", (1, 0), (1, 0), soft),
                    ("BOX", (0, 0), (-1, -1), 0.45, border),
                    ("LEFTPADDING", (0, 0), (0, 0), 0),
                    ("RIGHTPADDING", (0, 0), (0, 0), 0),
                    ("LEFTPADDING", (1, 0), (1, 0), 7),
                    ("RIGHTPADDING", (1, 0), (1, 0), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        return table

    def metric_table() -> Table:
        values = [
            (percent(summary["aiRate"]), "AI率", red),
            (percent(summary["suspiciousAiRate"]), "疑似AI率", amber),
            (percent(summary["humanRate"]), "人写作率", green),
            (f"[{summary['suspiciousChars']}]", "疑似AI写作字数", red),
            (f"[{summary['totalChars']}]", "总字数", blue_dark),
        ]
        value_cells = []
        label_cells = []
        for value, label, color in values:
            value_style = ParagraphStyle(f"AICheckMetricValue{label}", parent=metric_value_style, textColor=color)
            value_cells.append(Paragraph(safe(value), value_style))
            label_cells.append(Paragraph(safe(label), metric_label_style))
        table = Table([value_cells, label_cells], colWidths=[doc.width / 5] * 5, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                    ("BOX", (0, 0), (-1, -1), 0.45, border),
                    ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E5EAF1")),
                    ("TOPPADDING", (0, 0), (-1, 0), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 2),
                    ("TOPPADDING", (0, 1), (-1, 1), 0),
                    ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        return table

    detection_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    title = str(scan.get("documentTitle") or "untitled")
    author = str(scan.get("author") or "unknown")
    scope = "ChatGPT 讯飞星火 Gemini Kimi Claude 文心一言 通义千问 智谱AI 百川智能 360智脑 豆包 DeepSeek（包括但不限于）"
    meta_data = [
        [paragraph(f"检测时间：{detection_time}", meta_style), paragraph(f"检测文献：{title}", meta_style)],
        [paragraph(f"作者：{author}", meta_style), paragraph("检测类型：AIGC 写作检测", meta_style)],
        [paragraph(f"官网：{OFFICIAL_SITE}", meta_style), paragraph(f"检测范围：{scope}", meta_style)],
    ]
    meta_table = Table(meta_data, colWidths=[doc.width * 0.38, doc.width * 0.62], hAlign="LEFT")
    meta_table.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 2), (1, 2)),
                ("BACKGROUND", (0, 0), (-1, -1), softer),
                ("BOX", (0, 0), (-1, -1), 0.45, border),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E9EEF5")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    story: list[Any] = [
        Paragraph("AI-check / EasyIdea 论文AI率检测报告", title_style),
        Paragraph("独立 DOCX 解析生成 · 可选择文本 PDF · 结果仅供参考", subtitle_style),
        meta_table,
        Spacer(1, 6 * mm),
        section_title("检测结果"),
        Spacer(1, 3 * mm),
        metric_table(),
        Spacer(1, 6 * mm),
        section_title("疑似片段分布图"),
        Spacer(1, 3 * mm),
        DistributionBars([row["result"].aigcValue for row in rows]),
        Spacer(1, 6 * mm),
        section_title("原文内容 / 疑似AI写作率"),
        Spacer(1, 3 * mm),
    ]

    table_data: list[list[Any]] = [[paragraph("原文内容", table_header_style), paragraph("疑似AI写作率", table_header_style)]]
    for row in rows:
        result: Result = row["result"]
        reason = report_reason(result)
        value = max(0.0, min(1.0, result.aigcValue))
        original_cell: list[Any] = [paragraph(row["text"], body_style)]
        if reason:
            original_cell.extend([Spacer(1, 2), paragraph(f"检测原因：{reason}", reason_style)])
        score_color = color_for_value(value)
        score = Paragraph(f'<font color="{score_color.hexval()}">AIGC值：{value:.3f}</font>', score_style)
        table_data.append([original_cell, score])

    report_table = LongTable(table_data, colWidths=[doc.width * 0.74, doc.width * 0.26], repeatRows=1, hAlign="LEFT", splitByRow=1)
    table_style_commands: list[tuple[Any, ...]] = [
        ("BACKGROUND", (0, 0), (-1, 0), soft),
        ("BOX", (0, 0), (-1, -1), 0.45, border),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E5EAF1")),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    for row_index in range(1, len(table_data)):
        if row_index % 2 == 0:
            table_style_commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#FCFDFF")))
    report_table.setStyle(TableStyle(table_style_commands))
    story.append(report_table)
    story.extend(
        [
            Spacer(1, 7 * mm),
            section_title("说明"),
            Spacer(1, 2.5 * mm),
            paragraph("本报告由 AI-check 根据当前助手逐段判断结果生成，AIGC值与文章质量无关，仅供参考。", note_style),
            paragraph("本报告不代表知网、维普、Turnitin 或其他闭源平台检测结果。", note_style),
            paragraph("疑似AI生成段落中的“片段”为检测自动划分，与原文自然段可能不同。", note_style),
            paragraph(f"EasyIdea 官网：{OFFICIAL_SITE}", note_style),
        ]
    )

    def draw_page(canvas_obj: Any, doc_obj: Any) -> None:
        canvas_obj.saveState()
        footer_y = 10 * mm
        canvas_obj.setStrokeColor(colors.HexColor("#E5E7EB"))
        canvas_obj.setLineWidth(0.5)
        canvas_obj.line(doc_obj.leftMargin, footer_y + 4, page_w - doc_obj.rightMargin, footer_y + 4)
        canvas_obj.setFont(font_name, 7.5)
        canvas_obj.setFillColor(muted)
        canvas_obj.drawString(doc_obj.leftMargin, footer_y - 1, "AI-check / EasyIdea · 参考报告")
        canvas_obj.drawRightString(page_w - doc_obj.rightMargin, footer_y - 1, f"第 {doc_obj.page} 页")
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)


def generate_report(scan_path: Path, results_path: Path, out_dir: Path) -> dict[str, Any]:
    scan = load_json(scan_path)
    results = normalize_results(load_json(results_path))
    rows = get_report_rows(scan, results)
    summary = calculate_summary(scan, rows)
    out_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    html_path = out_dir / f"AI-check-EasyIdea-AI-rate-report-{timestamp}.html"
    pdf_path = out_dir / f"AI-check-EasyIdea-AI-rate-report-{timestamp}.pdf"
    summary_path = out_dir / f"AI-check-EasyIdea-AI-rate-summary-{timestamp}.json"

    html_path.write_text(render_html(scan, rows, summary), encoding="utf-8", newline="\n")
    pdf_error = ""
    try:
        render_pdf(pdf_path, scan, rows, summary)
    except Exception as exc:
        pdf_error = str(exc)
        pdf_path = Path("")

    output = {
        **summary,
        "htmlPath": str(html_path.resolve()),
        "pdfPath": str(pdf_path.resolve()) if str(pdf_path) else "",
        "pdfError": pdf_error,
        "reportPath": str(pdf_path.resolve()) if str(pdf_path) else str(html_path.resolve()),
        "easyIdea": OFFICIAL_SITE,
    }
    write_json(summary_path, output)
    output["summaryPath"] = str(summary_path.resolve())
    return output


def command_scan(args: argparse.Namespace) -> int:
    source = Path(args.docx)
    data = scan_docx(source, title=args.title, author=args.author, max_chars=args.max_chars)
    if args.out:
        write_json(Path(args.out), data)
    print(json.dumps(data if not args.out else {"scanPath": str(Path(args.out).resolve()), "chunks": len(data["chunks"]), "totalChars": data["totalChars"]}, ensure_ascii=False, indent=2))
    return 0


def command_report(args: argparse.Namespace) -> int:
    output = generate_report(Path(args.scan_json), Path(args.results_json), Path(args.out_dir))
    print(json.dumps(output, ensure_ascii=False, indent=2))
    return 0


def command_features(args: argparse.Namespace) -> int:
    text = args.text
    if args.file:
        text = Path(args.file).read_text(encoding="utf-8")
    print(json.dumps(asdict(scan_trace_features(text)), ensure_ascii=False, indent=2))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="AI-check deterministic DOCX scanner and report generator")
    sub = parser.add_subparsers(dest="command", required=True)

    scan = sub.add_parser("scan", help="read DOCX and write chunk/traceFeatures JSON")
    scan.add_argument("docx", help="source DOCX path")
    scan.add_argument("--out", help="output scan JSON path")
    scan.add_argument("--title", help="override document title")
    scan.add_argument("--author", help="override author")
    scan.add_argument("--max-chars", type=int, default=MAX_CHARS_PER_CHUNK, help="maximum Chinese-aware chars per chunk")
    scan.set_defaults(func=command_scan)

    report = sub.add_parser("report", help="render HTML/PDF report from scan JSON and assistant results JSON")
    report.add_argument("scan_json", help="scan JSON path")
    report.add_argument("results_json", help="assistant-provided results JSON path")
    report.add_argument("--out-dir", required=True, help="output directory")
    report.set_defaults(func=command_report)

    features = sub.add_parser("features", help="scan deterministic trace features for text")
    features.add_argument("text", nargs="?", default="", help="text to scan")
    features.add_argument("--file", help="UTF-8 text file to scan")
    features.set_defaults(func=command_features)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return int(args.func(args))
    except Exception as exc:
        print(f"AI-check error: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
